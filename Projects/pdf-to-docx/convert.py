import sys
import os
from pdf2docx import Converter
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

if len(sys.argv) != 3:
    print("Usage: python convert.py input.pdf output.docx")
    sys.exit(1)

pdf_path = sys.argv[1]
output_path = sys.argv[2]

def convert_with_pdf2docx(pdf_path, output_path):
    print(f"🚀 Mencoba konversi dengan pdf2docx...")
    try:
        cv = Converter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        print(f"✅ Konversi selesai (pdf2docx): {output_path}")
        return True
    except Exception as e:
        print(f"⚠️ pdf2docx gagal: {e}")
        return False

def convert_with_fitz(pdf_path, output_path):
    print(f"🧠 Fallback ke PyMuPDF (fitz)...")
    doc = fitz.open(pdf_path)
    word_doc = Document()

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        if text.strip():
            word_doc.add_paragraph(text)
        else:
            pix = page.get_pixmap()
            img_path = f"temp_page_{page_num}.png"
            pix.save(img_path)
            word_doc.add_picture(img_path, width=Inches(6))
            os.remove(img_path)

        word_doc.add_page_break()

    word_doc.save(output_path)
    print(f"✅ Konversi selesai (fitz): {output_path}")

# Jalankan pdf2docx dulu, kalau gagal → fallback ke fitz
if not convert_with_pdf2docx(pdf_path, output_path):
    convert_with_fitz(pdf_path, output_path)

sys.exit(0)
