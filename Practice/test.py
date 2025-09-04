from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_quote(doc, text, color):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('“' + text + '”')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = color

def create_biodata_doc(path="biodata_muhammad_fikri_fahrezi.docx"):
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("Biodata")
    t_run.bold = True
    t_run.font.size = Pt(28)

    # decorative line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    l_run = line.add_run("—" * 50)
    l_run.font.size = Pt(14)
    l_run.font.color.rgb = RGBColor(0x36, 0x75, 0xE6)  # blue shade

    doc.add_paragraph()

    # Biodata fields
    p1 = doc.add_paragraph()
    p1.add_run("Nama: ").bold = True
    p1.add_run("muhammad fikri fahrezi")

    p2 = doc.add_paragraph()
    p2.add_run("Prodi: ").bold = True
    p2.add_run("informatika")

    doc.add_paragraph()

    # Quotes header
    qh = doc.add_paragraph()
    qh.alignment = WD_ALIGN_PARAGRAPH.LEFT
    qh_run = qh.add_run("Kata-kata Quotes")
    qh_run.bold = True
    qh_run.font.size = Pt(14)

    quotes = [
        "Kegigihan adalah kunci untuk membuka pintu impian yang belum terlihat.",
        "Setiap langkah kecil hari ini membentuk cerita besar esok.",
        "Percaya diri adalah modal utama untuk menaklukkan ketakutan dan rintangan."
    ]
    colors = [
        RGBColor(0x1F, 0x7A, 0xD9),
        RGBColor(0xD0, 0x39, 0x39),
        RGBColor(0x2E, 0x7D, 0x32)
    ]

    for t, c in zip(quotes, colors):
        add_quote(doc, t, c)

    doc.save("/home/runner/work/tes-sshd/tes-sshd/Projects/ai-docx/tmp/hasil-667011e7-a11b-4351-bd55-3ad51b9ca6892.docx")

if __name__ == "__main__":
    create_biodata_doc()
